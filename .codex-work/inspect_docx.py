from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    document = Document(source)

    payload = {
        "source": str(source.resolve()),
        "paragraphs": [
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": paragraph.text,
            }
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip()
        ],
        "tables": [
            {
                "index": table_index,
                "rows": [
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ],
            }
            for table_index, table in enumerate(document.tables)
        ],
        "sections": len(document.sections),
        "inline_shapes": len(document.inline_shapes),
        "style_tokens": {},
        "section_tokens": [],
    }

    style_ids = {
        "Normal": "Normal",
        "Heading 1": "Heading1",
        "Heading 2": "Heading2",
        "Heading 3": "Heading3",
        "Heading 4": "Heading4",
    }
    for style_name, style_id in style_ids.items():
        style = next(style for style in document.styles if style.style_id == style_id)
        payload["style_tokens"][style_name] = {
            "font": style.font.name,
            "size_pt": style.font.size.pt if style.font.size else None,
            "bold": style.font.bold,
            "italic": style.font.italic,
            "color": str(style.font.color.rgb) if style.font.color and style.font.color.rgb else None,
            "space_before_pt": style.paragraph_format.space_before.pt if style.paragraph_format.space_before else None,
            "space_after_pt": style.paragraph_format.space_after.pt if style.paragraph_format.space_after else None,
            "line_spacing": style.paragraph_format.line_spacing,
            "keep_with_next": style.paragraph_format.keep_with_next,
        }

    for section in document.sections:
        payload["section_tokens"].append({
            "page_width_in": section.page_width.inches,
            "page_height_in": section.page_height.inches,
            "top_margin_in": section.top_margin.inches,
            "bottom_margin_in": section.bottom_margin.inches,
            "left_margin_in": section.left_margin.inches,
            "right_margin_in": section.right_margin.inches,
            "header_distance_in": section.header_distance.inches,
            "footer_distance_in": section.footer_distance.inches,
            "different_first_page": section.different_first_page_header_footer,
        })
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    if len(sys.argv) > 3:
        minimum = int(sys.argv[3])
        for paragraph in payload["paragraphs"]:
            if paragraph["index"] >= minimum:
                print(
                    f'{paragraph["index"]:03d} '
                    f'[{paragraph["style"]}] {paragraph["text"]}'
                )


if __name__ == "__main__":
    main()
