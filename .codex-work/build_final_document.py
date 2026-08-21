from __future__ import annotations

import base64
import copy
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor

from create_er_diagram import build_er_diagram


ROOT = Path(r"C:\VS Code\aut-ins-api")
SOURCE = ROOT / "docs" / "entregables" / "Proyecto_Aplicacion_AUT-INS_2_ETAPA_4_CERRADA.docx"
OUTPUT = ROOT / "docs" / "entregables" / "Proyecto_Aplicacion_AUT-INS_DOCUMENTACION_FINAL.docx"
LOGO_SVG = ROOT / "public" / "assets" / "img" / "aut-ins-logo.svg"
LOGO_PNG = ROOT / ".codex-work" / "aut-ins-logo.png"
ER_PNG = ROOT / ".codex-work" / "aut-ins-er-diagram.png"

BLACK = "000000"
WHITE = "FFFFFF"
LIGHT_GRAY = "F2F2F2"
NOTE_GRAY = "EDEDED"
MID_GRAY = "666666"
BLUE_1 = "365F91"
BLUE_2 = "4F81BD"
TABLE_WIDTH = 9648


def style_by_id(document: Document, style_id: str):
    return next(style for style in document.styles if style.style_id == style_id)


def configure_page_section(section, *, landscape: bool):
    """Apply the document's page geometry while preserving linked headers and footers."""
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True


def set_run_font(run, name: str = "Arial", size: float | None = None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_text(paragraph, text: str, *, bold=False, italic=False, color=None, size=10.5, align=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph()
    paragraph.style = style_by_id(document, f"Heading{level}")
    paragraph.add_run(text)
    return paragraph


def add_body(document: Document, text: str, *, bold_lead: str | None = None, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = document.add_paragraph()
    paragraph.style = style_by_id(document, "Normal")
    paragraph.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=10.5, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=10.5, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, italic=italic)
    return paragraph


def add_bullets(document: Document, items: list[str]):
    bullet_style = style_by_id(document, "ListBullet")
    for item in items:
        paragraph = document.add_paragraph(style=bullet_style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(item)
        set_run_font(run, size=10.5)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int]):
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"Los anchos deben sumar {TABLE_WIDTH}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def format_cell(cell, text: str, *, header=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    set_run_font(run, size=9.0, bold=header, color=WHITE if header else BLACK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if header:
        shade_cell(cell, BLACK)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int], *, caption: str | None = None, aligns: list | None = None):
    if caption:
        paragraph = document.add_paragraph()
        paragraph.style = style_by_id(document, "Normal")
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(caption)
        set_run_font(run, size=9.5, bold=True)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = style_by_id(document, "TableGrid")
    for index, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        format_cell(cell, header, header=True, align=(aligns[index] if aligns else WD_ALIGN_PARAGRAPH.CENTER))
    set_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        row = table.add_row()
        set_row_cant_split(row)
        for col_index, (cell, value) in enumerate(zip(row.cells, values)):
            align = aligns[col_index] if aligns else WD_ALIGN_PARAGRAPH.LEFT
            format_cell(cell, value, align=align)
            if row_index % 2 == 0:
                shade_cell(cell, LIGHT_GRAY)
    set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_note(document: Document, title: str, text: str):
    paragraph = document.add_paragraph()
    paragraph.style = style_by_id(document, "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), NOTE_GRAY)
    p_pr.append(shd)
    lead = paragraph.add_run(title)
    set_run_font(lead, size=9.5, bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, size=9.5, italic=True, color="444444")
    return paragraph


def add_code(document: Document, code: str):
    paragraph = document.add_paragraph()
    paragraph.style = style_by_id(document, "Normal")
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EFEFEF")
    p_pr.append(shd)
    run = paragraph.add_run(code)
    set_run_font(run, name="Consolas", size=8.3)
    return paragraph


def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE_2)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def delete_table_rows_except_header(table):
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)


def refill_existing_table(table, headers: list[str], rows: list[list[str]], widths: list[int], aligns=None):
    while len(table.columns) < len(headers):
        table.add_column(Inches(1))
    delete_table_rows_except_header(table)
    for index, header in enumerate(headers):
        format_cell(table.rows[0].cells[index], header, header=True, align=(aligns[index] if aligns else WD_ALIGN_PARAGRAPH.CENTER))
    set_repeat_header(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        row = table.add_row()
        set_row_cant_split(row)
        for col_index, value in enumerate(values):
            format_cell(row.cells[col_index], value, align=(aligns[col_index] if aligns else WD_ALIGN_PARAGRAPH.LEFT))
            if row_index % 2 == 0:
                shade_cell(row.cells[col_index], LIGHT_GRAY)
    set_table_geometry(table, widths)


def find_paragraph(document: Document, starts_with: str):
    return next(p for p in document.paragraphs if p.text.strip().startswith(starts_with))


def replace_paragraph(document: Document, starts_with: str, text: str, *, style_id: str | None = None):
    paragraph = find_paragraph(document, starts_with)
    if style_id:
        paragraph.style = style_by_id(document, style_id)
    set_paragraph_text(paragraph, text)
    return paragraph


def extract_logo():
    svg = LOGO_SVG.read_text(encoding="utf-8", errors="ignore")
    match = re.search(r"href=\"data:image/png;base64,([^\"]+)\"", svg)
    if not match:
        raise RuntimeError("No se encontró la imagen PNG incrustada en el logotipo SVG.")
    LOGO_PNG.write_bytes(base64.b64decode(match.group(1)))


def update_front_matter(document: Document):
    extract_logo()
    cover_cell = document.tables[0].cell(0, 1)
    cover_cell.text = ""
    cover_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cover_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(LOGO_PNG), width=Inches(1.12))

    index_rows = [
        ["Información general", "Equipo, funciones, objetivo y problemática."],
        ["Etapa 1", "Actores, UFP, cadena de valor y diagrama general."],
        ["Etapa 2", "Arquitectura, tecnologías y decisiones de solución."],
        ["Etapa 3", "Requisitos, modelo normalizado, migración y base de datos."],
        ["Etapa 4", "Recursos, rutas, contratos, validaciones, permisos y Postman."],
        ["Etapa 5", "Desarrollo realizado y estado técnico de la API."],
        ["Etapa 6", "Validaciones ejecutadas y estrategia de pruebas."],
        ["Etapa 7", "Front end demostrativo y experiencia por roles."],
        ["Etapa 8", "Evidencias, instalación y control de entrega."],
        ["Conclusiones", "Resultados, alcance real y continuidad del proyecto."],
        ["Anexos", "Glosario, comandos y trazabilidad Scrum."],
    ]
    refill_existing_table(
        document.tables[1],
        ["Apartado", "Contenido"],
        index_rows,
        [2500, 7148],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    team_rows = [
        ["1", "Carlos Eduardo Martínez Morales", "Scrum Master y desarrollador"],
        ["2", "Daen Sánchez Marín", "Líder técnico y desarrollador"],
        ["3", "Fernando Pérez Acuautla", "Desarrollador"],
        ["4", "Fernando Aguilar Velázquez", "Desarrollador"],
        ["5", "Pedro Jair Suárez Flores", "Responsable de pruebas"],
    ]
    refill_existing_table(
        document.tables[2],
        ["Núm.", "Integrante", "Responsabilidad"],
        team_rows,
        [800, 5000, 3848],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    actor_rows = [
        ["Aspirante", "Captura sus datos, selecciona carrera y entrega su expediente digital."],
        ["Aspirante", "Consulta el avance y sustituye solamente los documentos observados."],
        ["Coordinación de Admisiones", "Revisa la información y los documentos del aspirante."],
        ["Coordinación de Admisiones", "Emite dictamen de aceptación, rechazo u observación con motivo."],
        ["Coordinación de Admisiones", "Solicita la notificación externa correspondiente al aspirante."],
        ["Control Escolar", "Enrola al aspirante aceptado y genera su matrícula y acceso inicial."],
        ["Control Escolar", "Administra estados, datos y antecedentes de los alumnos."],
        ["Control Escolar", "Atiende los mensajes de carácter administrativo del alumno."],
        ["Coordinación Académica", "Administra grupos dentro de su carrera y controla su capacidad."],
        ["Coordinación Académica", "Asigna o cambia el grupo del alumno cuando exista disponibilidad."],
        ["Coordinación Académica", "Atiende mensajes académicos dirigidos a un coordinador específico."],
        ["Alumno", "Consulta su matrícula, estado, carrera, grupo, periodo y turno."],
        ["Alumno", "Envía mensajes a Control Escolar o a un coordinador específico."],
        ["Sistema", "Aplica integridad, historial, control de proceso activo y validación de cupo."],
        ["Sistema", "Protege el acceso, registra notificaciones y conserva trazabilidad."],
    ]
    refill_existing_table(
        document.tables[3],
        ["Actor", "Responsabilidad"],
        actor_rows,
        [3150, 6498],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    replacements = {
        "5. Asignación Académica de Grupo y Materias": "5. Asignación Académica de Grupo y Control de Cupo por Coordinación",
        "6. Consulta de Información Personal, Grupo": "6. Consulta de Información Personal, Estado e Inscripción",
    }
    for prefix, text in replacements.items():
        replace_paragraph(document, prefix, text, style_id="Heading3")

    ufp_updates = {
        8: [
            ["Elemento", "Información"],
            ["Nombre descriptivo", "UFP-05: Asignación Académica de Grupo y Control de Cupo"],
            ["Qué hace", "Permite a Coordinación administrar grupos de su carrera, definir capacidad y asignar alumnos únicamente cuando exista disponibilidad. El catálogo de materias queda fuera del núcleo de esta versión."],
            ["Actores involucrados", "Coordinación Académica, Sistema."],
            ["Problemática que resuelve", "Desorganización en la distribución de alumnos y saturación de grupos."],
        ],
        9: [
            ["Elemento", "Información"],
            ["Nombre descriptivo", "UFP-06: Consulta de Información Personal, Estado e Inscripción"],
            ["Qué hace", "Permite al alumno consultar su matrícula, estado, carrera, grupo, periodo y turno desde un portal privado."],
            ["Actores involucrados", "Alumno, Sistema."],
            ["Problemática que resuelve", "Reduce consultas presenciales y ofrece visibilidad inmediata de la situación escolar registrada."],
        ],
        10: [
            ["Elemento", "Información"],
            ["Nombre descriptivo", "UFP-07: Gestión de Solicitudes de Ayuda y Corrección de Datos"],
            ["Qué hace", "Mantiene conversaciones internas. Los mensajes administrativos llegan a la bandeja compartida de Control Escolar y los académicos pueden dirigirse a un coordinador específico."],
            ["Actores involucrados", "Alumno, Control Escolar, Coordinación Académica, Sistema."],
            ["Problemática que resuelve", "Evita solicitudes informales y conserva un historial consultable de la atención."],
        ],
        11: [
            ["Elemento", "Información"],
            ["Nombre descriptivo", "UFP-08: Historial, Indicadores y Control de Seguridad"],
            ["Qué hace", "Permite consultar procesos históricos por CURP y preparar indicadores de admisión e inscripción; la protección de acceso se define mediante autenticación y roles."],
            ["Actores involucrados", "Control Escolar, Sistema."],
            ["Problemática que resuelve", "Falta de trazabilidad y de información consolidada para el seguimiento administrativo."],
        ],
        14: [
            ["Elemento", "Información"],
            ["Nombre descriptivo", "UFP-SUG-03: Subsanación y Re-evaluación Documental"],
            ["Qué hace", "Permite que un aspirante con expediente observado acceda mediante un token temporal y sustituya únicamente los documentos indicados, sin crear otro folio ni repetir el registro."],
            ["Actores involucrados", "Aspirante, Coordinación de Admisiones, Sistema."],
            ["Problemática que resuelve", "Evita folios duplicados y conserva la versión anterior como evidencia."],
        ],
    }
    for table_index, values in ufp_updates.items():
        table = document.tables[table_index]
        for row_index, row_values in enumerate(values):
            for cell_index, value in enumerate(row_values):
                format_cell(table.rows[row_index].cells[cell_index], value, header=row_index == 0)
        set_table_geometry(table, [2500, 7148])

    replace_paragraph(
        document,
        "El equipo deberá decidir la arquitectura",
        "Se adoptó una arquitectura cliente-servidor con una API REST modular dentro de un monolito Node.js y Express. MySQL funciona como persistencia relacional mediante Prisma ORM. El mismo servidor publica actualmente una interfaz web estática de demostración; la integración de esa interfaz con los endpoints de negocio se realizará conforme avance la implementación.",
    )
    replace_paragraph(
        document,
        "Hasta este momento, las tecnologías confirmadas",
        "Las tecnologías se registran de acuerdo con su uso real o previsto. La documentación distingue los componentes ya implementados de aquellos cuyo contrato fue aprobado pero todavía requiere desarrollo.",
    )
    azure_paragraph = replace_paragraph(
        document,
        "Microsoft Azure se utilizará únicamente",
        "Microsoft Azure se reserva exclusivamente para alojar MySQL. El código se conserva en el repositorio privado de GitHub; el servicio de ejecución de la API podrá ser Render u otra plataforma compatible con Node.js, siempre que utilice variables de entorno y una conexión segura a la base alojada.",
    )
    azure_paragraph.paragraph_format.page_break_before = False

    technology_rows = [
        ["Node.js 22 y Express 5", "Servidor HTTP y API modular. Implementado.", "Permiten ejecutar TypeScript compilado, publicar el front y separar rutas por módulo dentro de una base sencilla de mantener."],
        ["TypeScript 7", "Lenguaje del backend. Implementado.", "Aporta tipado estático, facilita refactorizaciones y detecta inconsistencias antes de desplegar."],
        ["MySQL 8.4", "Persistencia relacional. Implementado localmente.", "Sus relaciones, restricciones e índices son adecuados para identidad, expedientes, estados, inscripciones e historial."],
        ["Prisma ORM 7.9", "Modelo, cliente, migración y datos semilla. Implementado.", "Centraliza el esquema, genera acceso tipado y permite reproducir la estructura mediante migraciones."],
        ["Microsoft Azure", "Alojamiento previsto de MySQL.", "Centralizará la base sin alojar el repositorio ni sustituir el servicio de ejecución de la API."],
        ["JWT", "Autenticación y permisos. Diseñado; pendiente en backend.", "El contrato permite sesiones sin estado en el servidor y separación de permisos por rol y carrera."],
        ["EmailJS", "Notificaciones externas al aspirante. Diseñado; pendiente de integración.", "Se limita a confirmaciones, observaciones y dictámenes del proceso de admisión."],
        ["CORS y dotenv", "Orígenes y variables de entorno. Implementados.", "Evitan valores sensibles en el código y permiten adaptar el servidor a distintos entornos."],
        ["HTML5, CSS3 y Bootstrap 5.3", "Interfaz responsive. Implementada como demostración.", "Ofrecen formularios y paneles claros sin introducir un framework de front adicional."],
        ["jsPDF 4.2", "PDF de acceso inicial. Implementado en la demostración.", "Genera localmente el comprobante con matrícula y contraseña temporal al concluir el enrolamiento demostrativo."],
        ["Git y GitHub", "Control de versiones y repositorio privado. Implementados.", "Mantienen historial, ramas y revisión del código sin publicar credenciales ni datos reales."],
        ["Postman", "Contrato y pruebas de API. Colección implementada.", "Agrupa 35 solicitudes y prepara la ejecución reproducible cuando los endpoints de negocio estén disponibles."],
    ]
    refill_existing_table(
        document.tables[16],
        ["Tecnología", "Uso dentro del proyecto", "Justificación"],
        technology_rows,
        [1750, 2700, 5198],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )


def trim_from_stage_three(document: Document):
    marker = find_paragraph(document, "3. Tercera etapa. Diseño de la base de datos")._element
    body = document._element.body
    start_index = list(body).index(marker)
    for element in list(body)[start_index:]:
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def clone_stage_four(document: Document, source_document: Document):
    source_body = source_document._element.body
    elements = list(source_body)
    start = elements.index(find_paragraph(source_document, "4. Cuarta etapa. Diseño de la API")._element)
    end = elements.index(find_paragraph(source_document, "5. Quinta etapa. Desarrollo")._element)
    destination_body = document._element.body
    for element in elements[start:end]:
        destination_body.insert(len(destination_body) - 1, copy.deepcopy(element))
    stage_heading = find_paragraph(document, "4. Cuarta etapa. Diseño de la API")
    stage_heading.paragraph_format.page_break_before = False


def add_stage_three(document: Document):
    document.add_page_break()
    add_heading(document, "3. Tercera etapa. Diseño e implementación de la base de datos", 1)
    add_note(
        document,
        "Estado de la etapa. ",
        "SCRUM-13 a SCRUM-17 se consideran concluidos para la línea base actual. El modelo fue normalizado, convertido en un esquema Prisma, materializado en una migración MySQL y acompañado por catálogos iniciales.",
    )
    add_body(document, "El alcance prioritario permanece en admisión e inscripción. No se incorporó una entidad de materias porque ese catálogo es complementario y no condiciona el flujo principal aprobado.")

    add_heading(document, "3.1 Propósito del diseño de la base de datos", 2)
    add_body(document, "Diseñar una persistencia relacional que registre cada intento de admisión, mantenga la identidad única por CURP, conserve expedientes y documentos históricos, permita el enrolamiento y la inscripción, controle cupos y soporte la comunicación institucional sin borrar información relevante.")

    add_heading(document, "3.2 Objetivos de la base de datos", 2)
    objectives = [
        ["Centralización", "Concentrar identidad, admisión, enrolamiento, inscripción e historial en una sola estructura relacional."],
        ["Integridad", "Aplicar claves, restricciones e índices que impidan referencias inválidas y duplicidades funcionales."],
        ["Historial por CURP", "Conservar varios procesos de una persona y permitir solamente uno activo a la vez."],
        ["Documentos", "Mantener versiones documentales y señalar una única versión vigente por tipo y expediente."],
        ["Seguridad", "Separar cuentas, roles y credenciales; almacenar únicamente hashes de contraseña."],
        ["Cupo", "Relacionar grupo, periodo e inscripción y calcular disponibilidad mediante estados que ocupan lugar."],
        ["Mensajería", "Conservar conversaciones, remitente real y lecturas por usuario."],
        ["Reproducibilidad", "Poder reconstruir la base mediante migraciones y datos semilla."],
    ]
    add_table(document, ["Objetivo", "Descripción"], objectives, [2200, 7448], caption="Tabla 3.1. Objetivos del modelo de datos.")

    add_heading(document, "3.3 Análisis de requisitos y entidades (SCRUM-13)", 2)
    add_heading(document, "3.3.1 Requisitos funcionales", 3)
    functional = [
        ["RF-01", "Registrar la identidad de la persona y crear un intento de admisión con carrera y datos de contacto.", "Aspirante / Sistema"],
        ["RF-02", "Generar folio y expediente; recibir documentos por tipo y versión.", "Sistema / Aspirante"],
        ["RF-03", "Revisar documentos, registrar observaciones y emitir dictamen.", "Admisiones"],
        ["RF-04", "Permitir la subsanación sin duplicar folio y conservar la versión sustituida.", "Aspirante / Admisiones"],
        ["RF-05", "Impedir dos procesos activos para una misma CURP y conservar procesos cerrados.", "Sistema"],
        ["RF-06", "Enrolar un expediente aceptado y generar una matrícula única.", "Control Escolar / Sistema"],
        ["RF-07", "Administrar estados de alumno y registrar su historial.", "Control Escolar"],
        ["RF-08", "Administrar carreras, periodos y grupos con capacidad máxima.", "Control Escolar / Coordinación"],
        ["RF-09", "Inscribir un alumno en un grupo compatible sin superar su cupo.", "Control Escolar / Coordinación"],
        ["RF-10", "Consultar perfil e inscripción propios desde el portal del alumno.", "Alumno"],
        ["RF-11", "Crear conversaciones y mensajes dirigidos a Control Escolar o a Coordinación.", "Alumno / Personal"],
        ["RF-12", "Registrar las notificaciones externas enviadas al aspirante.", "Sistema"],
        ["RF-13", "Consultar el historial relacionado con una CURP sin eliminar procesos anteriores.", "Control Escolar"],
    ]
    add_table(document, ["ID", "Requisito funcional", "Actor"], functional, [900, 6648, 2100], caption="Tabla 3.2. Requisitos funcionales consolidados.", aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(document, "3.3.2 Requisitos no funcionales", 3)
    non_functional = [
        ["RNF-01", "MySQL debe mantener integridad referencial y restringir el borrado de información histórica."],
        ["RNF-02", "Las contraseñas se conservarán como hashes; nunca se devolverán ni almacenarán en texto plano."],
        ["RNF-03", "El acceso definitivo se protegerá por JWT y permisos de rol."],
        ["RNF-04", "Las operaciones críticas se ejecutarán en transacciones para evitar estados parciales."],
        ["RNF-05", "El esquema, la migración y los catálogos deberán ser reproducibles desde el repositorio."],
        ["RNF-06", "La configuración y las credenciales se suministrarán mediante variables de entorno."],
        ["RNF-07", "La API utilizará respuestas JSON uniformes y no expondrá detalles internos de Prisma o MySQL."],
        ["RNF-08", "Los datos y archivos de demostración serán ficticios; no se publicarán secretos ni documentos reales."],
    ]
    add_table(document, ["ID", "Requisito no funcional"], non_functional, [1000, 8648], caption="Tabla 3.3. Requisitos no funcionales.")

    add_heading(document, "3.4 Reglas de datos e integridad", 2)
    rules = [
        ["RN-01", "La CURP es única en tbl_persona; los intentos históricos se registran en tbl_aspirante."],
        ["RN-02", "tbl_proceso_activo permite un solo proceso vigente por persona."],
        ["RN-03", "Cada aspirante tiene como máximo un expediente y cada expediente puede originar como máximo un alumno."],
        ["RN-04", "Solo un expediente aceptado puede continuar al enrolamiento."],
        ["RN-05", "La matrícula es única de forma global en tbl_usuario."],
        ["RN-06", "Cada versión documental es única por expediente, tipo y número de versión."],
        ["RN-07", "tbl_documento_vigente identifica una sola versión actual por expediente y tipo documental."],
        ["RN-08", "Cada grupo pertenece a una carrera y a un periodo."],
        ["RN-09", "Cada alumno tiene como máximo una inscripción por periodo."],
        ["RN-10", "El grupo y la inscripción deben pertenecer al mismo periodo mediante una relación compuesta."],
        ["RN-11", "El cupo se calcula con inscripciones cuyo estado tenga ocupa_cupo = true."],
        ["RN-12", "Los cambios de expediente, alumno e inscripción se registran en tablas históricas."],
        ["RN-13", "Los registros históricos utilizan ON DELETE RESTRICT y se cierran mediante estados."],
        ["RN-14", "Una conversación se dirige al departamento de Control Escolar o a un coordinador autorizado."],
        ["RN-15", "Cada lectura de mensaje es única por mensaje y usuario."],
    ]
    add_table(document, ["ID", "Regla de datos"], rules, [1000, 8648], caption="Tabla 3.4. Reglas consolidadas de integridad.")

    add_heading(document, "3.5 Modelo lógico definitivo", 2)
    entities = [
        ["Identidad", "tbl_persona", "Identidad única, CURP, nombre y contacto principal."],
        ["Identidad", "tbl_aspirante", "Cada intento de ingreso asociado con persona y carrera."],
        ["Identidad", "tbl_proceso_activo", "Apuntador único al intento vigente de una persona."],
        ["Admisiones", "tbl_expediente", "Folio, estado, observaciones y dictamen."],
        ["Admisiones", "tbl_cat_estado_expediente", "Estados, orden, bloqueo de CURP y condición terminal."],
        ["Documentos", "tbl_cat_tipo_documento", "Tipos institucionales, obligatoriedad y vigencia."],
        ["Documentos", "tbl_documento", "Metadatos, revisión y versiones del archivo."],
        ["Documentos", "tbl_documento_vigente", "Versión actual por expediente y tipo."],
        ["Seguridad", "tbl_usuario", "Matrícula, hash, tipo de usuario y estado de acceso."],
        ["Alumnos", "tbl_alumno", "Alumno originado por un expediente aceptado."],
        ["Alumnos", "tbl_cat_estado_alumno", "Activo, bajas y egreso con reglas de vigencia."],
        ["Personal", "tbl_empleado", "Cuenta del personal, rol y carrera opcional."],
        ["Personal", "tbl_cat_rol_empleado", "Admisiones, Control Escolar y Coordinación."],
        ["Oferta", "tbl_cat_carrera", "Carreras activas y sus relaciones."],
        ["Oferta", "tbl_cat_periodo", "Ciclo escolar, fechas y estado."],
        ["Oferta", "tbl_grupo", "Clave, turno, grado, capacidad, carrera y periodo."],
        ["Inscripción", "tbl_inscripcion", "Relación alumno-grupo-periodo y responsable."],
        ["Inscripción", "tbl_cat_estado_inscripcion", "Estados que indican si una inscripción ocupa cupo."],
        ["Historial", "tbl_historial_inscripcion", "Cambios de grupo y estado de inscripción."],
        ["Historial", "tbl_historial_estado_expediente", "Transiciones del expediente."],
        ["Historial", "tbl_historial_estado_alumno", "Transiciones académicas del alumno."],
        ["Mensajería", "tbl_conversacion", "Hilo, área destino, responsable y estado."],
        ["Mensajería", "tbl_mensaje", "Contenido, remitente real y fecha de envío."],
        ["Mensajería", "tbl_mensaje_lectura", "Lectura por mensaje y usuario."],
        ["Notificaciones", "tbl_notificacion_aspirante", "Evidencia del envío externo y su resultado."],
    ]
    add_table(document, ["Dominio", "Tabla física", "Responsabilidad"], entities, [1700, 3200, 4748], caption="Tabla 3.5. Catálogo de las 25 tablas implementadas.")

    add_heading(document, "3.5.1 Identidad e historial por CURP", 3)
    add_body(document, "La identidad se concentra en tbl_persona, donde la CURP es única. Cada nuevo intento se registra en tbl_aspirante, por lo que una persona puede conservar varios procesos históricos sin duplicar sus datos principales. tbl_proceso_activo funciona como autoridad para impedir dos procesos simultáneos.")
    add_heading(document, "3.5.2 Expediente y control documental", 3)
    add_body(document, "El expediente conserva el folio y el dictamen. Los archivos se modelan como registros independientes; cada sustitución crea una versión y tbl_documento_vigente señala la versión que debe mostrar la aplicación. La versión anterior permanece disponible como evidencia.")
    add_heading(document, "3.5.3 Cuenta de acceso y roles", 3)
    add_body(document, "tbl_usuario centraliza la matrícula y el hash de contraseña para alumnos y empleados. La especialización uno a uno evita matrículas duplicadas y permite que el rol administrativo y la carrera de Coordinación residan en entidades controladas.")
    add_heading(document, "3.5.4 Periodos, grupos e inscripciones", 3)
    add_body(document, "La inscripción referencia al alumno, el periodo, el grupo, el estado y el empleado responsable. Una clave foránea compuesta obliga a que grupo e inscripción compartan periodo. El límite de cupo se valida de forma transaccional al contar únicamente estados que ocupan lugar.")
    add_heading(document, "3.5.5 Mensajería y notificaciones", 3)
    add_body(document, "La mensajería interna separa conversaciones, mensajes y lecturas. EmailJS no reemplaza este módulo: las notificaciones externas al aspirante se registran aparte para conservar tipo, destinatario, estado e identificador externo.")

    add_heading(document, "3.6 Relaciones y cardinalidades", 2)
    relations = [
        ["tbl_persona", "1:N", "tbl_aspirante", "Una identidad puede realizar varios intentos históricos."],
        ["tbl_persona", "1:0..1", "tbl_proceso_activo", "Solo puede existir un intento vigente por persona."],
        ["tbl_aspirante", "1:0..1", "tbl_expediente", "Un intento origina como máximo un expediente."],
        ["tbl_expediente", "1:N", "tbl_documento", "El expediente contiene versiones documentales."],
        ["tbl_expediente", "1:0..1", "tbl_alumno", "Solo el proceso aceptado puede originar un alumno."],
        ["tbl_usuario", "1:0..1", "tbl_alumno / tbl_empleado", "Una cuenta se especializa en un solo tipo."],
        ["tbl_cat_carrera", "1:N", "aspirantes, alumnos, empleados y grupos", "La carrera se reutiliza sin duplicar texto."],
        ["tbl_cat_periodo", "1:N", "tbl_grupo", "Un periodo ofrece varios grupos."],
        ["tbl_alumno", "1:N", "tbl_inscripcion", "El alumno conserva su trayectoria por periodo."],
        ["tbl_grupo", "1:N", "tbl_inscripcion", "El grupo recibe alumnos hasta su capacidad."],
        ["tbl_inscripcion", "1:N", "tbl_historial_inscripcion", "Cada movimiento queda registrado."],
        ["tbl_alumno", "1:N", "tbl_conversacion", "El alumno puede iniciar varios hilos."],
        ["tbl_conversacion", "1:N", "tbl_mensaje", "Cada hilo contiene múltiples intervenciones."],
        ["tbl_mensaje", "N:M", "tbl_usuario", "tbl_mensaje_lectura resuelve la lectura por usuario."],
        ["tbl_aspirante", "1:N", "tbl_notificacion_aspirante", "Se conserva el historial de envíos externos."],
    ]
    add_table(document, ["Origen", "Cardinalidad", "Destino", "Interpretación"], relations, [2100, 1600, 2500, 3448], caption="Tabla 3.6. Relaciones principales del modelo.", aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(document, "3.7 Criterios de integridad y vigencia", 2)
    add_bullets(document, [
        "La baja temporal mantiene el proceso activo y conserva bloqueada la CURP.",
        "El rechazo, la cancelación, la baja definitiva y el egreso permiten liberar tbl_proceso_activo sin eliminar antecedentes.",
        "Los catálogos incluyen banderas como bloquea_curp, es_terminal u ocupa_cupo para evitar condiciones codificadas como texto libre.",
        "Las operaciones de historial utilizan llaves foráneas y fechas, no sobrescrituras destructivas.",
        "Las relaciones históricas se protegen principalmente con ON DELETE RESTRICT.",
    ])

    add_heading(document, "3.8 Casos de validación del modelo", 2)
    validation_cases = [
        ["CURP con proceso activo", "Bloqueado", "tbl_proceso_activo ya contiene a la persona."],
        ["Solicitud rechazada y nuevo intento", "Permitido", "Se elimina solo el apuntador activo; el expediente anterior permanece."],
        ["Baja temporal y nuevo intento", "Bloqueado", "La relación académica sigue vigente."],
        ["Baja definitiva y nueva carrera", "Permitido", "Se crea otro aspirante para la misma persona y se conserva todo el historial."],
        ["Grupo al límite", "Bloqueado", "La transacción cuenta inscripciones que ocupan cupo antes de insertar."],
        ["Documento observado", "Sustitución controlada", "Se inserta una nueva versión y se actualiza tbl_documento_vigente."],
    ]
    add_table(document, ["Caso", "Resultado", "Fundamento"], validation_cases, [2900, 1900, 4848], caption="Tabla 3.7. Casos de validación del modelo.")

    add_heading(document, "3.9 Trazabilidad entre requisitos y entidades", 2)
    traceability = [
        ["Admisión", "persona, aspirante, proceso_activo, expediente", "Identidad, folio, vigencia y dictamen."],
        ["Documentación", "tipo_documento, documento, documento_vigente", "Tipos, versiones y archivo actual."],
        ["Enrolamiento", "usuario, alumno, estado_alumno", "Matrícula y condición académica."],
        ["Control de acceso", "usuario, empleado, rol_empleado", "Credenciales y permisos."],
        ["Inscripción", "periodo, grupo, inscripción, estado_inscripción", "Asignación y cupo."],
        ["Historial", "tres tablas de historial", "Transiciones auditables."],
        ["Mensajería", "conversación, mensaje, mensaje_lectura", "Atención interna y lecturas."],
        ["Correo a aspirante", "notificación_aspirante", "Evidencia de EmailJS."],
    ]
    add_table(document, ["Proceso", "Entidades", "Cobertura"], traceability, [2100, 3900, 3648], caption="Tabla 3.8. Trazabilidad funcional del modelo.")

    add_heading(document, "3.10 Criterios del diseño físico", 2)
    add_bullets(document, [
        "Nombres físicos en snake_case con prefijo tbl_ y catálogos identificados como tbl_cat_*.",
        "Identificadores enteros sin signo y fechas con precisión de milisegundos cuando se requiere trazabilidad.",
        "Restricciones UNIQUE para CURP, folio, matrícula, claves de catálogo y combinaciones funcionales.",
        "Índices para nombres, estados, fechas, carrera, periodo, cupo, bandejas de mensajes e historial.",
        "El borrado físico se evita en registros con valor histórico; los estados expresan la vigencia.",
    ])

    landscape = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page_section(landscape, landscape=True)
    add_heading(document, "3.10.1 Diagrama entidad–relación de la base de datos", 3)
    add_body(
        document,
        "La figura representa las 25 tablas del modelo físico definido en Prisma. Para conservar su legibilidad, "
        "muestra las claves primarias, foráneas y únicas más relevantes, junto con las cardinalidades esenciales; "
        "los atributos completos, restricciones e índices permanecen documentados en el esquema y la migración SQL.",
    )
    figure = document.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = Pt(3)
    figure.paragraph_format.space_after = Pt(4)
    figure.add_run().add_picture(str(ER_PNG), width=Inches(9.45))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)
    caption_run = caption.add_run("Figura 3.1. Diagrama entidad–relación del modelo físico AUT-INS.")
    set_run_font(caption_run, size=9.0, bold=True)

    portrait = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page_section(portrait, landscape=False)

    add_heading(document, "3.11 Primera Forma Normal (SCRUM-14)", 2)
    add_body(document, "La Primera Forma Normal exige que cada atributo contenga un valor atómico y que no existan grupos repetidos dentro de una fila. El modelo inicial almacenaba documentos como columnas del expediente y combinaba varios datos de identidad o lectura en una sola entidad; esas estructuras fueron separadas.")
    first_nf = [
        ["Documentos como columnas del expediente", "tbl_documento registra una fila por archivo, tipo y versión."],
        ["Varios procesos dentro del mismo registro personal", "tbl_persona se separa de tbl_aspirante y cada intento ocupa una fila."],
        ["Estado o rol capturado como texto libre", "Se usan catálogos o enumeraciones con un solo valor por atributo."],
        ["Lista de mensajes dentro de una conversación", "Cada mensaje es una fila de tbl_mensaje."],
        ["Lectura global incrustada en mensaje", "tbl_mensaje_lectura registra una fila por mensaje y usuario."],
        ["Nombre completo sin estructura", "Nombre, apellido paterno y apellido materno se almacenan por separado."],
    ]
    add_table(document, ["Situación analizada", "Corrección aplicada"], first_nf, [4000, 5648], caption="Tabla 3.9. Aplicación de la Primera Forma Normal.")
    add_note(document, "Resultado de 1FN. ", "Todas las tablas poseen filas identificables, valores atómicos y ausencia de grupos repetidos. Los archivos reales permanecen fuera de las columnas relacionales; la base conserva únicamente sus metadatos y ubicación.")

    add_heading(document, "3.12 Segunda Forma Normal (SCRUM-15)", 2)
    add_body(document, "La Segunda Forma Normal requiere cumplir 1FN y que cada atributo no clave dependa de la clave completa. La mayoría de las tablas utiliza una clave primaria simple, por lo que no puede presentar dependencias parciales. La revisión se concentró en las claves compuestas y combinaciones únicas.")
    second_nf = [
        ["tbl_documento_vigente", "(id_expediente, id_tipo_documento)", "id_documento depende de ambos valores: identifica la versión vigente de ese tipo dentro de ese expediente."],
        ["tbl_mensaje_lectura", "(id_mensaje, id_usuario)", "leido_en corresponde a la lectura de un mensaje por un usuario específico."],
        ["tbl_documento", "UNIQUE expediente + tipo + versión", "Los atributos del archivo dependen de la fila documental, no solo del expediente o del tipo."],
        ["tbl_grupo", "UNIQUE periodo + carrera + clave", "Los datos del grupo se conservan en tbl_grupo; carrera y periodo se resuelven mediante sus catálogos."],
        ["tbl_inscripcion", "UNIQUE alumno + periodo", "La fila representa la inscripción de ese alumno durante un periodo; el grupo compatible se valida por relación compuesta."],
    ]
    add_table(document, ["Tabla", "Clave revisada", "Comprobación"], second_nf, [2600, 3000, 4048], caption="Tabla 3.10. Revisión de dependencias completas en 2FN.")
    add_note(document, "Resultado de 2FN. ", "No existen atributos que dependan únicamente de una parte de una clave compuesta. Los datos descriptivos de tipos, estados, carreras y periodos residen en sus propias tablas.")

    add_heading(document, "3.13 Tercera Forma Normal (SCRUM-16)", 2)
    add_body(document, "La Tercera Forma Normal exige cumplir 2FN y eliminar dependencias transitivas entre atributos no clave. Se separaron las descripciones reutilizables y las responsabilidades que podrían depender indirectamente de otra entidad.")
    third_nf = [
        ["CURP y datos personales repetidos en cada trámite", "tbl_persona", "El aspirante referencia la identidad; no duplica sus atributos."],
        ["Carrera escrita en aspirante, alumno, empleado y grupo", "tbl_cat_carrera", "Cada entidad conserva únicamente id_carrera."],
        ["Contraseña en alumno y empleado", "tbl_usuario", "La cuenta concentra matrícula, hash, tipo y vigencia."],
        ["Descripción y reglas del estado en cada expediente", "tbl_cat_estado_expediente", "El expediente guarda id_estado_expediente."],
        ["Descripción del estado académico", "tbl_cat_estado_alumno", "El alumno guarda id_estado_alumno."],
        ["Regla de ocupación de cupo repetida", "tbl_cat_estado_inscripcion", "La bandera ocupa_cupo pertenece al estado, no a cada inscripción."],
        ["Datos del archivo actual repetidos", "tbl_documento_vigente", "El apuntador vigente evita duplicar metadatos."],
        ["Nombre del remitente dentro del mensaje", "tbl_usuario", "tbl_mensaje conserva id_usuario_remitente."],
    ]
    add_table(document, ["Dependencia detectada", "Entidad separada", "Resultado"], third_nf, [3200, 2700, 3748], caption="Tabla 3.11. Eliminación de dependencias transitivas en 3FN.")
    add_note(document, "Resultado de 3FN. ", "Las descripciones y reglas compartidas dependen de sus propias claves. El esquema evita que un cambio de carrera, rol o estado obligue a actualizar textos repetidos en múltiples tablas.")

    add_heading(document, "3.14 Diseño e implementación de la base de datos (SCRUM-17)", 2)
    add_heading(document, "3.14.1 Archivos implementados", 3)
    files = [
        ["prisma/schema.prisma", "25 modelos, 8 enumeraciones, relaciones, restricciones e índices."],
        ["prisma/migrations/20260809000000_init/migration.sql", "Migración inicial que crea las tablas y claves foráneas de MySQL."],
        ["prisma/seed.ts", "Carga idempotente de estados, roles y tipos documentales."],
        ["prisma.config.ts", "Configuración de Prisma 7 y lectura segura de DATABASE_URL."],
        ["src/config/database.ts", "Cliente Prisma compartido mediante el adaptador MariaDB/MySQL."],
        ["scripts/local-mysql.ps1", "Preparación de una instancia local aislada para desarrollo."],
    ]
    add_table(document, ["Archivo", "Responsabilidad"], files, [3900, 5748], caption="Tabla 3.12. Evidencias de implementación de la base.")

    add_heading(document, "3.14.2 Catálogos iniciales", 3)
    catalogs = [
        ["Estado de expediente", "PENDIENTE_DOCUMENTOS, PENDIENTE_REVISION, OBSERVADO, ACEPTADO, RECHAZADO, CANCELADO"],
        ["Estado de alumno", "ACTIVO, BAJA_TEMPORAL, BAJA_DEFINITIVA, EGRESADO"],
        ["Estado de inscripción", "ACTIVA, CANCELADA, FINALIZADA"],
        ["Roles", "ADMISIONES, CONTROL_ESCOLAR, COORDINACION"],
        ["Tipos documentales", "Acta, certificado, identificación, comprobante, CURP y fotografía"],
    ]
    add_table(document, ["Catálogo", "Valores semilla"], catalogs, [2600, 7048], caption="Tabla 3.13. Catálogos cargados de forma idempotente.")

    add_heading(document, "3.14.3 Operaciones transaccionales", 3)
    add_bullets(document, [
        "Pre-registro: persona, aspirante, expediente y proceso activo.",
        "Dictamen: estado, historial, notificación y liberación del proceso cuando sea terminal.",
        "Enrolamiento: usuario, alumno e historial inicial.",
        "Inscripción y cambio de grupo: bloqueo, validación de carrera y periodo, conteo de cupo y movimiento histórico.",
        "Sustitución documental: nueva versión, actualización del apuntador vigente y conservación de la anterior.",
        "Baja definitiva o egreso: estado, historial y liberación controlada del proceso activo.",
    ])

    add_heading(document, "3.14.4 Consultas de referencia", 3)
    add_body(document, "Consultar el proceso vigente de una CURP:")
    add_code(document, "SELECT p.curp, e.folio, ce.codigo AS estado\nFROM tbl_persona p\nJOIN tbl_proceso_activo pa ON pa.id_persona = p.id_persona\nJOIN tbl_aspirante a ON a.id_aspirante = pa.id_aspirante\nJOIN tbl_expediente e ON e.id_aspirante = a.id_aspirante\nJOIN tbl_cat_estado_expediente ce ON ce.id_estado_expediente = e.id_estado_expediente\nWHERE p.curp = ?;")
    add_body(document, "Consultar el historial completo de una CURP:")
    add_code(document, "SELECT a.fecha_registro, e.folio, c.nombre AS carrera, ce.codigo AS estado\nFROM tbl_persona p\nJOIN tbl_aspirante a ON a.id_persona = p.id_persona\nJOIN tbl_expediente e ON e.id_aspirante = a.id_aspirante\nJOIN tbl_cat_carrera c ON c.id_carrera = a.id_carrera\nJOIN tbl_cat_estado_expediente ce ON ce.id_estado_expediente = e.id_estado_expediente\nWHERE p.curp = ? ORDER BY a.fecha_registro;")
    add_body(document, "Calcular cupo ocupado:")
    add_code(document, "SELECT g.capacidad_maxima, COUNT(i.id_inscripcion) AS ocupados\nFROM tbl_grupo g\nLEFT JOIN tbl_inscripcion i ON i.id_grupo = g.id_grupo\nLEFT JOIN tbl_cat_estado_inscripcion ei\n  ON ei.id_estado_inscripcion = i.id_estado_inscripcion AND ei.ocupa_cupo = TRUE\nWHERE g.id_grupo = ? GROUP BY g.id_grupo, g.capacidad_maxima;")

    add_heading(document, "3.14.5 Validación y ejecución", 3)
    add_code(document, "npm run mysql:setup\nnpm run db:validate\nnpm run db:deploy\nnpm run db:seed\nnpm run db:studio")
    add_note(document, "Resultado de SCRUM-17. ", "El esquema fue validado con Prisma, la migración forma parte del repositorio y los catálogos pueden cargarse nuevamente sin crear duplicados. La conexión productiva a Azure deberá suministrarse mediante DATABASE_URL; no se incluyen credenciales en el documento ni en Git.")


def add_stage_four_status(document: Document):
    add_heading(document, "4.7 Correspondencia entre el diseño y la implementación actual", 2)
    add_body(document, "Las 35 rutas de los apartados anteriores representan el contrato aprobado y la colección inicial de Postman. Al cierre de esta revisión, el servidor Express publica la ruta de diagnóstico y monta los prefijos modulares, pero los routers de negocio todavía no contienen controladores operativos. Esta distinción evita presentar como ejecutado lo que aún es diseño.")
    current = [
        ["GET /", "Implementado", "Entrega public/index.html y los recursos del front demostrativo."],
        ["GET /api/health", "Implementado", "Devuelve estado, versión, fecha y ocho módulos registrados."],
        ["/api/auth, /aspirantes, /admisiones, /control-escolar", "Montados", "Existen routers modulares; sus endpoints de negocio siguen pendientes."],
        ["/api/coordinacion, /alumnos, /mensajes, /historial", "Montados", "Existen routers modulares; sus endpoints de negocio siguen pendientes."],
        ["Rutas no definidas", "Implementado", "El middleware notFound devuelve una respuesta JSON 404 uniforme."],
    ]
    add_table(document, ["Superficie", "Estado", "Comportamiento comprobado"], current, [3350, 1500, 4798], caption="Tabla 4.12. Estado real de las rutas en el código.")


def add_stage_five(document: Document):
    add_heading(document, "5. Quinta etapa. Desarrollo", 1)
    add_note(document, "Estado de la etapa. ", "Desarrollo parcial controlado. La base ejecutable, la persistencia y la demostración web están disponibles; la lógica de los endpoints de negocio, JWT y EmailJS aún debe implementarse para considerar terminado el backend funcional.")

    add_heading(document, "5.1 Objetivo y criterio de implementación", 2)
    add_body(document, "El desarrollo se organizó por módulos para que cada área avance sin redefinir el modelo de datos ni invadir responsabilidades. El servidor, la base y los contratos constituyen una línea base estable; cualquier cambio de requisito debe registrarse antes de modificar el esquema o la colección.")

    add_heading(document, "5.2 Arquitectura del repositorio", 2)
    architecture = [
        ["src/app.ts", "Configura Express, CORS, límites JSON, archivos estáticos, /api y manejo de errores."],
        ["src/server.ts", "Inicia el servicio y realiza un cierre controlado ante SIGINT o SIGTERM."],
        ["src/config", "Centraliza variables de entorno y el cliente Prisma."],
        ["src/modules", "Separa auth, aspirantes, admisiones, Control Escolar, Coordinación, alumnos, mensajes, historial y common."],
        ["prisma", "Contiene esquema, migración, semilla y documentación de persistencia."],
        ["public", "Contiene la interfaz web demostrativa y sus recursos."],
        ["docs/postman", "Contiene la colección inicial con los contratos de la API."],
        ["scripts", "Incluye verificadores, generación de Postman y preparación local de MySQL."],
    ]
    add_table(document, ["Ubicación", "Responsabilidad"], architecture, [3000, 6648], caption="Tabla 5.1. Estructura técnica implementada.")

    add_heading(document, "5.3 Módulos y límites funcionales", 2)
    modules = [
        ["auth", "Login, renovación, cierre y permisos JWT.", "Router creado; lógica pendiente."],
        ["aspirantes", "Pre-registro, consulta y subsanación.", "Router creado; lógica pendiente."],
        ["admisiones", "Bandeja, documentos y dictamen.", "Router creado; lógica pendiente."],
        ["control-escolar", "Enrolamiento, matrícula, altas, bajas y correcciones.", "Router creado; lógica pendiente."],
        ["coordinacion", "Grupos, capacidad y asignaciones por carrera.", "Router creado; lógica pendiente."],
        ["alumnos", "Perfil, estado e inscripción propios.", "Router creado; lógica pendiente."],
        ["mensajes", "Conversaciones, envío y lecturas.", "Router creado; lógica pendiente."],
        ["historial", "Procesos históricos por CURP.", "Router creado; lógica pendiente."],
        ["common", "Salud, errores y rutas inexistentes.", "Implementado y operativo."],
    ]
    add_table(document, ["Módulo", "Alcance", "Estado actual"], modules, [1800, 4700, 3148], caption="Tabla 5.2. Estado de los módulos del backend.")

    add_heading(document, "5.4 Componentes desarrollados", 2)
    components = [
        ["Servidor", "Express 5, CORS configurable, límites de entrada, publicación estática y cierre controlado.", "Operativo"],
        ["Diagnóstico", "GET /api/health con versión y listado de módulos.", "Operativo"],
        ["Errores", "Respuesta uniforme para AppError, errores internos y rutas inexistentes.", "Operativo"],
        ["Persistencia", "Cliente Prisma, 25 tablas, migración inicial y datos semilla.", "Implementado"],
        ["Contratos", "Colección Postman v2.1 con 35 solicitudes y once UFP.", "Implementado"],
        ["Front", "Landing, solicitud, login y paneles de cuatro roles con datos de demostración.", "Demostración funcional"],
        ["JWT", "Diseño de claims, vigencia y permisos.", "Pendiente en backend"],
        ["EmailJS", "Alcance y tipos de notificación definidos.", "Pendiente de integración"],
        ["Almacenamiento documental", "Modelo de metadatos y versionado definido.", "Proveedor de archivos pendiente"],
    ]
    add_table(document, ["Componente", "Resultado", "Estado"], components, [1900, 5700, 2048], caption="Tabla 5.3. Resultado del desarrollo actual.")

    add_heading(document, "5.5 Configuración y comandos", 2)
    commands = [
        ["npm install", "Instala dependencias y genera el cliente Prisma."],
        ["npm run mysql:setup", "Prepara MySQL local aislado en el puerto 3307 y genera .env."],
        ["npm run db:deploy", "Aplica las migraciones existentes."],
        ["npm run db:seed", "Carga o actualiza los catálogos iniciales."],
        ["npm run dev", "Ejecuta el servidor con recarga durante el desarrollo."],
        ["npm run check", "Valida esquema, tipos, estructura, Postman y compilación estática."],
        ["npm run build / npm start", "Compila y ejecuta dist/server.js."],
    ]
    add_table(document, ["Comando", "Propósito"], commands, [3000, 6648], caption="Tabla 5.4. Comandos de desarrollo y ejecución.")

    add_heading(document, "5.6 Consideraciones de despliegue", 2)
    add_bullets(document, [
        "Render puede ejecutar el servicio Node.js mediante npm run build y npm start.",
        "Azure alojará únicamente MySQL y deberá permitir conexiones desde el servicio de la API.",
        "DATABASE_URL, CORS_ORIGIN y cualquier secreto deben definirse como variables privadas del entorno.",
        "En producción se ejecuta prisma migrate deploy; migrate dev queda reservado al desarrollo.",
        "La salud del servicio se comprueba mediante GET /api/health.",
    ])
    add_note(document, "Límite actual. ", "La compatibilidad de despliegue está preparada a nivel de scripts y servidor, pero no existe evidencia de una instancia productiva publicada ni de una conexión Azure configurada. El documento no la presenta como completada.")


def add_stage_six(document: Document):
    add_heading(document, "6. Sexta etapa. Pruebas", 1)
    add_body(document, "Responsable registrado: Pedro Jair Suárez Flores. Esta revisión documenta las comprobaciones disponibles al 12 de agosto de 2026 y separa las validaciones ejecutadas de los casos que requieren la implementación futura de los endpoints de negocio.")

    add_heading(document, "6.1 Resumen de validaciones ejecutadas", 2)
    checks = [
        ["Prisma", "npm run db:validate", "Esquema válido.", "Conforme"],
        ["Tipos de base", "npm run db:typecheck", "Configuración y semilla sin errores de TypeScript.", "Conforme"],
        ["Estructura", "npm run check:structure", "9 módulos y 17 archivos TypeScript reconocidos.", "Conforme"],
        ["Postman", "npm run check:postman", "35 rutas únicas y 11 UFP cubiertas.", "Conforme"],
        ["Tipos del backend", "npm run typecheck", "Sin errores de tipado.", "Conforme"],
        ["Compilación", "npm run build", "dist generado correctamente.", "Conforme"],
    ]
    add_table(document, ["Área", "Ejecución", "Resultado observado", "Estado"], checks, [1500, 2450, 4148, 1550], caption="Tabla 6.1. Validaciones automatizadas ejecutadas.", aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

    add_heading(document, "6.2 Pruebas de humo HTTP", 2)
    smoke = [
        ["GET /api/health", "200", "application/json", "Servicio y módulos disponibles."],
        ["GET /", "200", "text/html", "Interfaz principal entregada por Express."],
        ["GET /api/aspirantes", "404", "application/json", "Resultado esperado: router sin endpoint implementado."],
        ["GET /api/ruta-inexistente", "404", "application/json", "Middleware de ruta inexistente operativo."],
    ]
    add_table(document, ["Solicitud", "HTTP", "Contenido", "Comprobación"], smoke, [2700, 900, 2100, 3948], caption="Tabla 6.2. Resultados de las pruebas de humo locales.", aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(document, "6.3 Revisión manual del front", 2)
    front_checks = [
        ["Navegación pública", "Landing, solicitud e inicio de sesión visibles y navegables.", "Conforme"],
        ["Admisiones", "Bandeja de solicitudes, filtros y revisión de expediente disponibles.", "Conforme en demostración"],
        ["Control Escolar", "Enrolamiento, alumnos, cambio de estado, acceso y PDF disponibles.", "Conforme en demostración"],
        ["Coordinación", "Grupos, capacidad, asignación y bandeja personal disponibles.", "Conforme en demostración"],
        ["Alumno", "Perfil, inscripción y caja de mensajes con selección de destinatario disponibles.", "Conforme en demostración"],
        ["Sesión", "Configuración de 10 segundos cerró la sesión por inactividad y regresó al login.", "Conforme"],
        ["Consola", "No se detectaron errores durante el recorrido realizado.", "Conforme"],
    ]
    add_table(document, ["Escenario", "Evidencia observada", "Resultado"], front_checks, [2200, 5548, 1900], caption="Tabla 6.3. Resumen de prueba manual del front.")

    add_heading(document, "6.4 Alcance de la colección Postman", 2)
    add_body(document, "La colección incluye aserciones iniciales de código HTTP, formato JSON, tiempo de respuesta y captura de identificadores. Su estructura fue validada automáticamente; sin embargo, la ejecución de los 35 flujos contra la API real permanece pendiente porque los routers de negocio aún no implementan esos contratos.")

    add_heading(document, "6.5 Casos funcionales pendientes", 2)
    add_bullets(document, [
        "Login válido, inválido, expirado y acceso con rol incorrecto mediante JWT real.",
        "Pre-registro correcto, formatos inválidos y bloqueo por proceso activo de la CURP.",
        "Carga, observación y sustitución documental con almacenamiento real.",
        "Dictamen aceptado, rechazado y observado con historial y EmailJS.",
        "Enrolamiento transaccional, matrícula única y cambio obligatorio de contraseña.",
        "Inscripción con cupo, sobrecupo, grupo de otra carrera y periodo cerrado.",
        "Mensajería autorizada, conversación ajena, lectura y cierre.",
        "Consultas históricas por CURP y aislamiento de datos por rol y carrera.",
    ])
    add_note(document, "Conclusión de pruebas. ", "La línea base compila, el esquema y los artefactos documentales son consistentes, y la demostración web completa el recorrido visual. No puede declararse aprobada la API de negocio hasta ejecutar los casos funcionales anteriores contra controladores y una base de pruebas.")


def add_stage_seven(document: Document):
    add_heading(document, "7. Séptima etapa. Front end", 1)
    add_note(document, "Estado de la etapa. ", "Demostración funcional concluida; integración con la API real pendiente. Responsable previsto: Carlos Eduardo Martínez Morales.")

    add_heading(document, "7.1 Propósito y enfoque visual", 2)
    add_body(document, "La interfaz presenta un tema claro con fondo gris azulado, acentos institucionales, tarjetas y paneles responsivos. Bootstrap se utiliza para la estructura y los componentes, mientras que CSS propio unifica la identidad AUT-INS y adapta las vistas a escritorio y móvil.")

    add_heading(document, "7.2 Vistas y funciones disponibles", 2)
    views = [
        ["Inicio", "Explica el flujo, ofrece acceso a solicitud y login, y presenta las cuatro experiencias principales."],
        ["Solicitud", "Captura datos, CURP, carrera, contacto, consentimiento y documentación de demostración."],
        ["Login", "Ofrece accesos discretos de prueba y redirige según el rol identificado."],
        ["Admisiones", "Métricas, búsqueda, filtro, expediente, documentos, observaciones y dictamen demostrativo."],
        ["Control Escolar", "Enrolamiento, matrícula, contraseña temporal, PDF, alumnos, estados y mensajes."],
        ["Coordinación", "Catálogo de grupos, capacidad, activación, asignación compatible y mensajes directos."],
        ["Alumno", "Matrícula, estado, carrera, grupo, periodo, turno y caja de mensajes."],
    ]
    add_table(document, ["Vista", "Alcance demostrativo"], views, [2200, 7448], caption="Tabla 7.1. Componentes del front implementado.")

    add_heading(document, "7.3 Sesión y separación por roles", 2)
    add_body(document, "La sesión demostrativa se guarda en sessionStorage. Si el usuario intenta abrir un panel sin sesión, regresa al login; si intenta entrar al panel de otro rol, se redirige al panel autorizado. El tiempo de inactividad es configurable entre 10 segundos y 30 minutos para facilitar la demostración, con 15 minutos como valor recomendado.")
    add_note(document, "Importante. ", "Este comportamiento simula la sesión en el navegador. La seguridad definitiva depende de JWT, validación de expiración en el servidor, hashes de contraseña y autorización en cada endpoint.")

    add_heading(document, "7.4 Expediente y enrolamiento demostrativo", 2)
    add_body(document, "Admisiones puede abrir los datos y la relación de documentos enviados por cada aspirante. Control Escolar recibe los expedientes aceptados y puede generar una matrícula y una contraseña temporal. El front permite descargar un PDF confidencial de acceso mediante jsPDF y solicita que la credencial se cambie en el primer ingreso.")

    add_heading(document, "7.5 Mensajería interna", 2)
    add_body(document, "El alumno elige entre el departamento de Control Escolar y un coordinador específico. Los mensajes dirigidos al departamento aparecen para los usuarios que comparten ese rol; los mensajes a Coordinación se muestran únicamente en la bandeja personal del coordinador seleccionado.")

    add_heading(document, "7.6 Persistencia demostrativa y límites", 2)
    add_bullets(document, [
        "Los datos de prueba se guardan localmente en el navegador y pueden restablecerse.",
        "Los documentos cargados se representan mediante metadatos; el almacenamiento real aún no está conectado.",
        "Las credenciales visibles son exclusivas del entorno de demostración y no deben reutilizarse.",
        "Las acciones del front no modifican todavía MySQL porque los endpoints de negocio siguen pendientes.",
        "EmailJS no se ejecuta desde esta demostración; solamente está definido el flujo esperado.",
    ])

    add_heading(document, "7.7 Criterios para la integración definitiva", 2)
    integration = [
        ["Autenticación", "Reemplazar demo-store por POST /api/auth/login y almacenar el JWT conforme a la estrategia de seguridad."],
        ["Datos", "Sustituir arreglos de demostración por llamadas a los endpoints autorizados."],
        ["Archivos", "Integrar carga multipart y previsualización desde un proveedor de almacenamiento seguro."],
        ["Errores", "Mostrar códigos y mensajes uniformes de la API sin exponer detalles internos."],
        ["Sesión", "Coordinar la inactividad del cliente con exp e invalidación del JWT."],
        ["Pruebas", "Ejecutar los mismos recorridos con datos aislados en un entorno de pruebas."],
    ]
    add_table(document, ["Área", "Trabajo de integración"], integration, [2200, 7448], caption="Tabla 7.2. Pendientes para conectar el front con la API.")


def add_stage_eight(document: Document):
    add_heading(document, "8. Octava etapa. Evidencias y entrega", 1)
    add_body(document, "La entrega actual reúne el código fuente, el modelo físico, la migración, los datos semilla, la colección de Postman, la documentación Markdown, la interfaz demostrativa y este documento maestro.")

    add_heading(document, "8.1 Inventario de evidencias", 2)
    evidence = [
        ["Código fuente", "src/, public/, scripts/", "Repositorio privado de GitHub"],
        ["Base de datos", "prisma/schema.prisma y migrations/", "Modelo y migración reproducibles"],
        ["Catálogos", "prisma/seed.ts", "Carga idempotente"],
        ["Diseño de API", "docs/postman/AUT-INS_API.postman_collection.json", "35 solicitudes y 11 UFP"],
        ["Documentación técnica", "docs/*.md y README.md", "Alcance, tecnología, datos, API y Scrum"],
        ["Front", "public/index.html y public/assets/", "Demostración por roles"],
        ["Documento académico", "Proyecto_Aplicacion_AUT-INS_DOCUMENTACION_FINAL.docx", "Memoria integral del proyecto"],
    ]
    add_table(document, ["Evidencia", "Ubicación", "Contenido"], evidence, [2100, 4748, 2800], caption="Tabla 8.1. Inventario de la entrega.")

    add_heading(document, "8.2 Instalación resumida", 2)
    add_code(document, "git clone https://github.com/jr-0205/aut-ins-api.git\ncd aut-ins-api\nnpm install\nnpm run mysql:setup\nnpm run db:deploy\nnpm run db:seed\nnpm run dev")
    add_body(document, "Después de iniciar el servidor, la interfaz está disponible en http://localhost:3000 y el diagnóstico en http://localhost:3000/api/health.")

    add_heading(document, "8.3 Repositorio y control de acceso", 2)
    paragraph = add_body(document, "Repositorio del proyecto: ", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_hyperlink(paragraph, "https://github.com/jr-0205/aut-ins-api", "Repositorio privado AUT-INS en GitHub")
    add_body(document, "El repositorio es privado. Solamente los colaboradores invitados pueden consultar el avance. No deben publicarse archivos .env, tokens, claves de Azure, documentos personales ni contraseñas de alumnos.")

    add_heading(document, "8.4 Estado general de cierre documental", 2)
    status = [
        ["Etapas 1 y 2", "Concluidas", "Actores, UFP, cadena de valor y decisiones tecnológicas consolidadas."],
        ["Etapa 3", "Concluida", "Modelo normalizado, esquema, migración y catálogos implementados."],
        ["Etapa 4", "Concluida", "Contrato de API y colección inicial revisados."],
        ["Etapa 5", "Parcial", "Base técnica operativa; endpoints de negocio pendientes."],
        ["Etapa 6", "Parcial", "Validaciones técnicas y front ejecutados; pruebas funcionales de API pendientes."],
        ["Etapa 7", "Demostración concluida", "Interfaz por roles disponible; integración real pendiente."],
        ["Etapa 8", "Concluida para esta entrega", "Evidencias e instalación reunidas."],
    ]
    add_table(document, ["Etapa", "Estado", "Justificación"], status, [2000, 2200, 5448], caption="Tabla 8.2. Estado verificable del proyecto.")


def add_conclusions_and_appendices(document: Document):
    add_heading(document, "Conclusiones", 1)
    add_body(document, "AUT-INS cuenta con un análisis funcional consistente, un modelo relacional normalizado y reproducible, un contrato de API amplio, una base modular compilable y una interfaz demostrativa que permite recorrer el proceso por roles. La estructura resuelve la principal preocupación del proyecto: conservar la identidad e historial de una CURP sin permitir procesos vigentes duplicados.")
    add_body(document, "La revisión también establece un límite importante: la existencia de 35 rutas en Postman no significa que los controladores estén terminados. El estado real del backend se documenta como base ejecutable con diagnóstico, persistencia y routers modulares, mientras que el front opera con datos locales de demostración. Esta separación permite continuar el desarrollo sin generar expectativas incorrectas.")
    add_body(document, "El siguiente incremento debe concentrarse en implementar el flujo completo de admisión e inscripción de extremo a extremo: autenticación, pre-registro, expediente, dictamen, enrolamiento, inscripción, mensajes e historial. Después deberán ejecutarse las pruebas funcionales de Postman y conectar el front a la API.")

    add_heading(document, "Anexo A. Glosario", 1)
    glossary = [
        ["Aspirante", "Persona que inicia un intento de admisión y todavía no tiene una cuenta de alumno activa."],
        ["Expediente", "Conjunto lógico de folio, estado, dictamen y documentos de un intento de admisión."],
        ["Enrolamiento", "Conversión controlada de un expediente aceptado en alumno y cuenta institucional."],
        ["Proceso activo", "Único intento vigente asociado con una persona; bloquea otro registro con la misma CURP."],
        ["UFP", "Unidad Funcional de Proceso que agrupa una capacidad del sistema."],
        ["JWT", "Token firmado que identifica al usuario y limita la vigencia de la sesión."],
        ["ORM", "Capa que representa las tablas y relaciones mediante modelos de código."],
        ["Migración", "Script versionado que reproduce una modificación del esquema físico."],
        ["Dato semilla", "Registro inicial controlado, como un rol o estado de catálogo."],
    ]
    add_table(document, ["Término", "Definición"], glossary, [2200, 7448], caption="Tabla A.1. Glosario del proyecto.")

    add_heading(document, "Anexo B. Comandos de verificación", 1)
    add_code(document, "npm run db:validate\nnpm run db:typecheck\nnpm run check:structure\nnpm run check:postman\nnpm run typecheck\nnpm run build\nnpm start")
    add_body(document, "La prueba mínima de disponibilidad se realiza con GET /api/health. Para pruebas completas se deberá importar docs/postman/AUT-INS_API.postman_collection.json y configurar baseUrl, tokens e identificadores del entorno de pruebas.")

    add_heading(document, "Anexo C. Trazabilidad Scrum", 1)
    scrum = [
        ["SCRUM-5", "Identificación de actores", "Integrado en 1.1"],
        ["SCRUM-6", "Matriz de actores y responsabilidades", "Integrado en 1.2"],
        ["SCRUM-7", "Unidades Funcionales de Proceso", "Integrado en 1.3 y 1.4"],
        ["SCRUM-8", "Cadena de valor", "Integrado en 1.5"],
        ["SCRUM-9", "Diseño de solución y tecnologías", "Integrado en etapa 2"],
        ["SCRUM-13", "Análisis de requisitos y entidades", "Integrado en 3.3 a 3.9"],
        ["SCRUM-14", "Primera Forma Normal", "Integrado en 3.11"],
        ["SCRUM-15", "Segunda Forma Normal", "Integrado en 3.12"],
        ["SCRUM-16", "Tercera Forma Normal", "Integrado en 3.13"],
        ["SCRUM-17", "Diseño e implementación de base", "Integrado en 3.14"],
        ["SCRUM-19 a 24", "Diseño y revisión de API", "Integrado en etapa 4"],
        ["Desarrollo", "Base técnica y front demostrativo", "Integrado en etapas 5 y 7"],
        ["Pruebas", "Validaciones ejecutadas y plan pendiente", "Integrado en etapa 6"],
    ]
    add_table(document, ["Tarea", "Entregable", "Ubicación"], scrum, [1900, 4600, 3148], caption="Tabla C.1. Relación entre el tablero y el documento.")


def set_image_alt_text(document: Document):
    descriptions = [
        ("Logo AUT-INS", "Logotipo monocromático del proyecto AUT-INS."),
        (
            "Diagrama general de actores y UFP",
            "Figura 1.1. Diagrama general de actores y Unidades Funcionales de Proceso, conservado del documento original.",
        ),
        (
            "Diagrama entidad–relación AUT-INS",
            "Figura 3.1. Modelo físico de 25 tablas organizado por flujo principal, referencias, documentos, seguridad, historial y mensajería.",
        ),
    ]
    for shape, (title, description) in zip(document.inline_shapes, descriptions):
        doc_properties = shape._inline.docPr
        doc_properties.set("title", title)
        doc_properties.set("descr", description)


def set_document_properties(document: Document):
    properties = document.core_properties
    properties.title = "AUT-INS - Documentación final del proyecto"
    properties.subject = "Análisis, diseño, base de datos, API, desarrollo, pruebas y front end"
    properties.author = "Equipo AUT-INS"
    properties.keywords = "AUT-INS, Scrum, API REST, MySQL, Prisma, inscripciones"
    properties.comments = "Versión consolidada a partir del documento académico y del repositorio de desarrollo."
    properties.modified = datetime.now()
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_er_diagram(ER_PNG)
    document = Document(SOURCE)
    source_document = Document(SOURCE)
    update_front_matter(document)
    trim_from_stage_three(document)
    add_stage_three(document)
    clone_stage_four(document, source_document)
    add_stage_four_status(document)
    add_stage_five(document)
    add_stage_six(document)
    add_stage_seven(document)
    add_stage_eight(document)
    add_conclusions_and_appendices(document)
    set_image_alt_text(document)
    set_document_properties(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
